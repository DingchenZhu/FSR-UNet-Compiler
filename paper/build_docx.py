"""Convert paper markdown files to a single DOCX document."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = Path(__file__).parent.parent / "docs"
OUTPUT = Path(__file__).parent / "论文_TVM编译器前端设计.docx"

# Files in final order
FILES = [
    "paper_abstract.md",
    "paper_chapter_1.md",
    "paper_background.md",
    "paper_chapters_3_4_5.md",
    "paper_chapter_6.md",
    "paper_chapter_7.md",
    "paper_references.md",
]

# ── Style setup ──────────────────────────────────────────────────────────────

def setup_styles(doc):
    styles = doc.styles

    # Normal
    n = styles["Normal"]
    n.font.name = "宋体"
    n.font.size = Pt(12)
    n._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = Pt(20)

    # Heading 1
    _set_heading(styles["Heading 1"], 18, bold=True, color="1F3864")
    # Heading 2
    _set_heading(styles["Heading 2"], 15, bold=True, color="2E4057")
    # Heading 3
    _set_heading(styles["Heading 3"], 13, bold=True, color="404040")
    # Heading 4
    _set_heading(styles["Heading 4"], 12, bold=True, color="555555")

    # Code block style (create if absent)
    if "Code" not in [s.name for s in styles]:
        code_style = styles.add_style("Code", 1)  # 1 = paragraph style
    else:
        code_style = styles["Code"]
    code_style.base_style = styles["Normal"]
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(9)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.paragraph_format.space_before = Pt(2)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.left_indent = Cm(0.5)
    # Light gray background via direct XML shading
    pPr = code_style.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def _set_heading(style, size, bold=True, color="000000"):
    style.font.name = "黑体"
    style.font.size = Pt(size)
    style.font.bold = bold
    r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
    style.font.color.rgb = RGBColor(r, g, b)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.keep_with_next = True


# ── Inline formatting ─────────────────────────────────────────────────────────

def add_inline(para, text):
    """Add text with inline markdown formatting (bold, inline code)."""
    # Split on **bold** and `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        else:
            if part:
                para.add_run(part)


# ── Table parsing ─────────────────────────────────────────────────────────────

def parse_table(doc, lines, start):
    """Parse a markdown table starting at `start`. Returns next line index."""
    rows = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        # Skip separator rows (e.g. |---|---|)
        if re.match(r"^\|[-| :]+\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1

    if not rows:
        return i

    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.clear()
            add_inline(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph()  # spacing after table
    return i


# ── Main parser ───────────────────────────────────────────────────────────────

def parse_markdown(doc, text):
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal in_code, code_buf
        if code_buf:
            for cl in code_buf:
                doc.add_paragraph(cl, style="Code")
        code_buf = []
        in_code = False

    while i < len(lines):
        line = lines[i]

        # ── Code fence ──
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Skip YAML front-matter and file-level metadata comments
        if stripped.startswith("---") or stripped.startswith("> "):
            i += 1
            continue

        # ── Blank line ──
        if not stripped:
            i += 1
            continue

        # ── Headings ──
        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            doc.add_heading(heading_text, level=level)
            i += 1
            continue

        # ── Table ──
        if stripped.startswith("|"):
            i = parse_table(doc, lines, i)
            continue

        # ── Horizontal rule ──
        if re.match(r"^[-*_]{3,}$", stripped):
            doc.add_paragraph("─" * 40, style="Normal")
            i += 1
            continue

        # ── List items ──
        m_li = re.match(r"^[-*]\s+(.*)", stripped)
        if m_li:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, m_li.group(1))
            i += 1
            continue

        m_li_num = re.match(r"^\d+\.\s+(.*)", stripped)
        if m_li_num:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, m_li_num.group(1))
            i += 1
            continue

        # ── Normal paragraph (may span multiple lines) ──
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_stripped = lines[i].strip()
            # Stop at blank line, heading, fence, or table
            if (not next_stripped or
                    next_stripped.startswith("#") or
                    next_stripped.startswith("```") or
                    next_stripped.startswith("|") or
                    next_stripped.startswith("---")):
                break
            para_lines.append(next_stripped)
            i += 1

        full_text = " ".join(para_lines)
        p = doc.add_paragraph(style="Normal")
        add_inline(p, full_text)

    if in_code:
        flush_code()


# ── Page number footer ────────────────────────────────────────────────────────

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    setup_styles(doc)
    add_page_numbers(doc)

    for filename in FILES:
        path = DOCS_DIR / filename
        if not path.exists():
            print(f"WARNING: {path} not found, skipping", file=sys.stderr)
            continue
        text = path.read_text(encoding="utf-8")
        parse_markdown(doc, text)
        # Page break between major files
        doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
